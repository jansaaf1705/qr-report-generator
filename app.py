from flask import Flask, render_template_string, request, send_file
import os
import uuid
import qrcode
import fitz
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from werkzeug.utils import secure_filename

app = Flask(__name__)

# Memory stores for mapping unique IDs to generated paths
REPORTS = {}
QR_CODES = {}

REPORT_FOLDER = "reports"
UPLOAD_FOLDER = "uploads"
QR_FOLDER = "qr"
ALLOWED_EXTENSIONS = {"pdf", "docx"}

app.config["REPORT_FOLDER"] = REPORT_FOLDER
app.config["UPLOAD_FOLDER"] = UPLOAD_FOLDER
app.config["QR_FOLDER"] = QR_FOLDER

os.makedirs(REPORT_FOLDER, exist_ok=True)
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(QR_FOLDER, exist_ok=True)


def allowed_file(filename):
    return "." in filename and filename.rsplit(".", 1)[1].lower() in ALLOWED_EXTENSIONS


def generate_qr(filename):
    unique_id = str(uuid.uuid4())
    
    # DYNAMIC FIX: Ensures the QR code points to your production URL on Render
    domain = request.host_url.rstrip('/')
    report_url = f"{domain}/report/{unique_id}"
    
    img = qrcode.make(report_url)
    qr_path = os.path.join(app.config["QR_FOLDER"], f"{unique_id}.png")
    img.save(qr_path)
    return unique_id, qr_path


def add_qr_to_pdf(pdf_path, qr_path, output_path):
    doc = fitz.open(pdf_path)
    page = doc[0]
    qr_size = 90
    margin = 20

    rect = fitz.Rect(
        page.rect.width - qr_size - margin,
        margin,
        page.rect.width - margin,
        margin + qr_size
    )

    page.insert_image(rect, filename=qr_path)
    doc.save(output_path)
    doc.close()


def add_qr_to_docx(docx_path, qr_path, output_path):
    doc = Document(docx_path)
    section = doc.sections[0]
    header = section.header
    
    if header.paragraphs:
        paragraph = header.paragraphs[0]
    else:
        paragraph = header.add_paragraph()

    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    run.add_picture(qr_path, width=Inches(1))
    doc.save(output_path)


HTML = """
<!DOCTYPE html>
<html>
<head>
    <title>QR Report Generator</title>
    <script src="https://cdn.tailwindcss.com"></script>
</head>
<body class="bg-gray-100">
<div class="min-h-screen flex items-center justify-center p-6">
    <div class="bg-white p-10 rounded-xl shadow-xl w-full max-w-xl">
        <h1 class="text-3xl font-bold text-center text-blue-700">
            QR Report Generator
        </h1>
        <p class="text-center text-gray-500 mt-3">
            Upload PDF or DOCX Report
        </p>
        
        <form method="POST" enctype="multipart/form-data" class="mt-8">
            <input
                type="file"
                name="report"
                accept=".pdf,.docx"
                required
                class="w-full border rounded-lg p-3">
            <button
                type="submit"
                class="w-full mt-5 bg-blue-700 text-white p-3 rounded-lg hover:bg-blue-800 transition">
                Upload & Embed QR
            </button>
        </form>

        {% if message %}
        <div class="mt-6 p-4 rounded bg-blue-50 text-blue-800 text-sm break-all">
            {{ message }}
        </div>
        {% endif %}

        {% if download_ready %}
        <div class="mt-6 border border-green-200 rounded-xl bg-green-50 p-6">
            <p class="text-green-800 font-bold text-center mb-4 text-base">🎉 File & QR Code Processed!</p>
            
            <div class="space-y-4">
                <div class="p-3 bg-white rounded-lg border border-gray-200">
                    <p class="text-xs font-semibold text-gray-500 mb-2 tracking-wide uppercase">📄 Report File Options</p>
                    <div class="flex gap-2">
                        <a href="/report/{{ report_id }}" target="_blank" 
                           class="flex-1 text-center bg-green-600 hover:bg-green-700 text-white py-2 rounded-lg text-sm font-medium transition">
                           View Report
                        </a>
                        <a href="/report/{{ report_id }}?download=true" 
                           class="flex-1 text-center bg-gray-800 hover:bg-gray-900 text-white py-2 rounded-lg text-sm font-medium transition">
                           Download Report
                        </a>
                    </div>
                </div>

                <div class="p-3 bg-white rounded-lg border border-gray-200">
                    <p class="text-xs font-semibold text-gray-500 mb-2 tracking-wide uppercase">🖼️ QR Code Image Options</p>
                    <div class="flex gap-2">
                        <a href="/qr/{{ report_id }}" target="_blank" 
                           class="flex-1 text-center bg-teal-600 hover:bg-teal-700 text-white py-2 rounded-lg text-sm font-medium transition">
                           View QR Code
                        </a>
                        <a href="/qr/{{ report_id }}?download=true" 
                           class="flex-1 text-center bg-gray-800 hover:bg-gray-900 text-white py-2 rounded-lg text-sm font-medium transition">
                           Download QR (.png)
                        </a>
                    </div>
                </div>
            </div>
        </div>
        {% endif %}
    </div>
</div>
</body>
</html>
"""


@app.route("/", methods=["GET", "POST"])
def home():
    message = ""
    download_ready = False
    report_id = ""

    if request.method == "POST":
        if "report" not in request.files:
            message = "No file selected."
        else:
            file = request.files["report"]

            if file.filename == "":
                message = "Please choose a file."
            elif allowed_file(file.filename):
                filename = secure_filename(file.filename)
                filepath = os.path.join(app.config["UPLOAD_FOLDER"], filename)
                file.save(filepath)

                report_id, qr_path = generate_qr(filename)
                
                output_report = os.path.join(
                    app.config["REPORT_FOLDER"],
                    f"QR_{filename}"
                )

                if filename.lower().endswith(".pdf"):
                    add_qr_to_pdf(filepath, qr_path, output_report)
                elif filename.lower().endswith(".docx"):
                    add_qr_to_docx(filepath, qr_path, output_report)

                REPORTS[report_id] = output_report
                QR_CODES[report_id] = qr_path
                
                download_ready = True
                message = f"Process verified. Unique ID: {report_id}"
            else:
                message = "Only PDF and DOCX files are allowed."

    return render_template_string(
        HTML, 
        message=message, 
        download_ready=download_ready, 
        report_id=report_id
    )


@app.route("/report/<report_id>")
def open_report(report_id):
    if report_id not in REPORTS:
        return "Report Not Found", 404
    
    should_download = request.args.get("download", "false").lower() == "true"
    return send_file(
        REPORTS[report_id], 
        as_attachment=should_download,
        download_name=os.path.basename(REPORTS[report_id])
    )


@app.route("/qr/<report_id>")
def open_qr(report_id):
    if report_id not in QR_CODES:
        return "QR Code Not Found", 404
    
    should_download = request.args.get("download", "false").lower() == "true"
    return send_file(
        QR_CODES[report_id], 
        as_attachment=should_download,
        download_name=f"QR_{report_id}.png"
    )


if __name__ == "__main__":
    # DYNAMIC FIX: Allows production cloud systems to explicitly define network endpoints
    port = int(os.environ.get("PORT", 5000))
    app.run(host="0.0.0.0", port=port)
